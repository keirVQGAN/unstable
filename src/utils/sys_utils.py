import os
import shutil
import pandas as pd
from pathlib import Path
from zipfile import ZipFile
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
import logging

def sort_files(directory):
    for path in Path(directory).iterdir():
        if path.is_file():
            prefix = path.stem.split('_')[0]
            destination = Path(directory) / prefix
            destination.mkdir(exist_ok=True)
            path.rename(destination / path.name)
        elif path.is_dir():
            sort_files(path)

def load_env_file(env_file_path, print_details=True):
    shutil.copy(env_file_path, '.env')
    load_dotenv()

    ENV_VARS = [
        'OPENAI_API_KEY',
        'STABLE_API_KEY',
        'UCARE_API_KEY_PUBLIC',
        'UCARE_API_KEY_SECRET',
        'OUT_PATH',
        'IN_PATH',
        'CONFIG_PATH'
    ]

    env_values = {var: os.getenv(var) for var in ENV_VARS}

    if print_details:
        for var, value in env_values.items():
            print(f"{var}: {value}")

    return env_values

def str2list(s, ignore_commas=False):
    if ignore_commas:
        return [s]
    if s.replace(',','').replace('.','').isdigit():
        return [str(i.strip()) for i in s.split(',')]
    return [s]

def create_dirs(dirs):
    if isinstance(dirs, str):
        dirs = [dirs]

    for path in dirs:
        Path(path).mkdir(parents=True, exist_ok=True)

def zip_folder(folder_path, output_filename):
    with ZipFile(output_filename, 'w') as zipf:
        for root, _, files in os.walk(folder_path):
            for file in files:
                zipf.write(
                    os.path.join(root, file),
                    arcname=os.path.relpath(
                        os.path.join(root, file),
                        folder_path
                    )
                )

def read_excel(path):
    df = pd.read_excel(path)
    return df.iloc[:, 0].tolist()

def read_docx(path):
    doc = Document(path)
    return [p.text for p in doc.paragraphs if p.text]

def read_txt(path):
    with open(path, 'r') as file:
        return file.readlines()

def write_docx(path, content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(path)

def write_txt_or_csv(path, content):
    with open(path, 'w') as file:
        file.write(content + '\n')

def write_xlsx(path, content):
    df = pd.DataFrame(content.split('\n'), columns=['Reference'])
    df.to_excel(path, index=False)

def setup_logging(log_file):
    logging.basicConfig(filename=log_file, level=logging.DEBUG)

def calculate_token_cost(token_count, cost_per_token):
    return token_count / 1000 * cost_per_token
